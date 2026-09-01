from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from .text_cleaning import clean_text, compact_spaces


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".csv", ".json"}


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str


def discover_documents(source_dir: Path) -> list[Path]:
    if not source_dir.exists():
        return []

    return sorted(
        path
        for path in source_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        text = load_pdf(path)
    elif suffix == ".docx":
        text = load_docx(path)
    elif suffix == ".xlsx":
        text = load_xlsx(path)
    elif suffix == ".csv":
        text = load_csv(path)
    elif suffix == ".json":
        text = load_json(path)
    else:
        raise ValueError(f"Unsupported file type: {path}")

    return LoadedDocument(path=path, text=clean_text(text))


def load_documents(paths: Iterable[Path]) -> list[LoadedDocument]:
    documents: list[LoadedDocument] = []
    for path in paths:
        document = load_document(path)
        if document.text.strip():
            documents.append(document)
    return documents


def load_pdf(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PDF loading needs PyMuPDF. Run: python -m pip install PyMuPDF") from exc

    parts: list[str] = []
    with fitz.open(path) as pdf:
        for index, page in enumerate(pdf, start=1):
            page_text = page.get_text("text")
            if not page_text.strip() or len(page_text.strip()) < 20:
                page_text = ocr_page(page, index)
            if page_text.strip():
                parts.append(f"[Page {index}]\n{page_text}")
    return "\n\n".join(parts)


def ocr_page(page: object, index: int) -> str:
    """Try OCR on a scanned page. Falls back to a note if tesseract is unavailable."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return f"[Page {index}] (OCR unavailable: install pytesseract and Pillow)"

    try:
        pixmap = page.get_pixmap(dpi=200)
        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        return pytesseract.image_to_string(image)
    except Exception:
        return f"[Page {index}] (OCR failed)"


def load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX loading needs python-docx. Run: python -m pip install python-docx") from exc

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                table_rows.append(" | ".join(values))

    return "\n".join(paragraphs + table_rows)


def load_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX loading needs openpyxl. Run: python -m pip install openpyxl") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [compact_spaces(str(value)) for value in row if value is not None and str(value).strip()]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def load_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            values = [compact_spaces(value) for value in row if value.strip()]
            if values:
                rows.append(" | ".join(values))
    return "\n".join(rows)


def load_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8"))
    return flatten_json(data)


def flatten_json(value: object, prefix: str = "") -> str:
    lines: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            next_prefix = f"{prefix}.{key}" if prefix else str(key)
            lines.append(flatten_json(item, next_prefix))
    elif isinstance(value, list):
        for index, item in enumerate(value):
            next_prefix = f"{prefix}[{index}]"
            lines.append(flatten_json(item, next_prefix))
    elif value is not None:
        label = f"{prefix}: " if prefix else ""
        lines.append(f"{label}{value}")

    return "\n".join(line for line in lines if line.strip())

